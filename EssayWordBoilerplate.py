import docx
import os
import calendar

from docx.shared import Pt
from datetime import datetime


# Generates the word document given the parameters passed in
def generateBoilerplate(student, professor, class_name, title, date):
    runValue = {
        0: student,
        1: professor,
        2: class_name,
        3: date,
        4: title
    }

    doc = docx.Document()

    paragraphs = []

    for i in range(5):
        paragraphs.append(doc.add_paragraph())

    # Set up the style for each of the paragraphs
    for i, paragraph in enumerate(paragraphs):
        format = paragraph.paragraph_format
        format.space_before = Pt(0)
        format.space_after = Pt(0)
        format.line_spacing = 2

        if i == 4:
            format.alignment = 1

        run = paragraph.add_run(runValue[i])
        font = run.font

        font.size = Pt(12)
        font.name = 'Times New Roman'

    # Save the Word file in the documents folder
    documents_folder = os.path.expanduser('~/Documents/')
    file = input("What do you want to call this .docx file? ")

    path = documents_folder + file + '.docx'

    doc.save(path)


# Get the MLA formatted current date
def getDateString():
    now = datetime.now()

    day = now.date().day
    month = calendar.month_name[now.date().month]
    year = now.date().year

    date = '{day} {month} {year}'.format(day=day, month=month, year=year)

    return date


# Main function to get inputs to pass to boilerplate
def main():
    student = input('What is your full name? ')
    professor = 'Professor ' + input('What is your professor\'s last name? ')
    class_name = input('What is the class name (preferably abbreviated version)? ')
    title = input('What is the title of the paper? ')
    date = getDateString()

    generateBoilerplate(student, professor, class_name, title, date)


if __name__ == '__main__':
    main()